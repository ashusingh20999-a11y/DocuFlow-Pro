from pathlib import Path
import pandas as pd, fitz, subprocess, tempfile, zipfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pypdf import PdfReader, PdfWriter
from PIL import Image
from pdf2docx import Converter

def convert_file(files,target,outdir):
    outdir.mkdir(parents=True,exist_ok=True); target=target.lower()
    if target=="merge_pdf":
        out=outdir/"merged.pdf"; w=PdfWriter()
        for p in files:
            if p.suffix.lower()!=".pdf": raise ValueError("Merge PDF requires PDF files only")
            for page in PdfReader(str(p)).pages: w.add_page(page)
        with out.open("wb") as f:w.write(f)
        return out
    if target=="split_pdf":
        if len(files)!=1 or files[0].suffix.lower()!=".pdf": raise ValueError("Select one PDF for split")
        r=PdfReader(str(files[0])); folder=outdir/"split_pages"; folder.mkdir()
        for i,page in enumerate(r.pages,1):
            w=PdfWriter(); w.add_page(page)
            with (folder/f"page_{i}.pdf").open("wb") as f:w.write(f)
        zpath=outdir/"split_pages.zip"
        with zipfile.ZipFile(zpath,"w",zipfile.ZIP_DEFLATED) as z:
            for p in folder.glob("*.pdf"): z.write(p,p.name)
        return zpath
    if target=="image_pdf":
        imgs=[Image.open(p).convert("RGB") for p in files]
        if not imgs: raise ValueError("No image selected")
        out=outdir/"images.pdf"; imgs[0].save(out,save_all=True,append_images=imgs[1:]); return out
    src=files[0]; ext=src.suffix.lower()
    if target=="docx" and ext==".pdf": return pdf_docx(src,outdir/"converted.docx")
    if target=="xlsx" and ext==".pdf": return pdf_xlsx(src,outdir/"converted.xlsx")
    if target=="pdf" and ext==".docx": return docx_pdf(src,outdir/"converted.pdf")
    if target=="pdf" and ext in [".xlsx",".xls",".csv"]: return table_pdf(src,outdir/"converted.pdf")
    if target=="xlsx" and ext==".csv":
        o=outdir/"converted.xlsx"; pd.read_csv(src).to_excel(o,index=False); return o
    if target=="csv" and ext in [".xlsx",".xls"]:
        o=outdir/"converted.csv"; pd.read_excel(src).to_csv(o,index=False); return o
    if target=="jpg" and ext==".pdf": return pdf_images(src,outdir)
    raise ValueError(f"Unsupported conversion: {ext} -> {target}")

def pdf_docx(src,out):
    converter=Converter(str(src))
    try: converter.convert(str(out),start=0,end=None)
    finally: converter.close()
    return out

def pdf_xlsx(src,out):
    rows=[]; pdf=fitz.open(src)
    try:
        for no,page in enumerate(pdf,1):
            tabs=page.find_tables()
            if tabs.tables:
                for t in tabs.tables: rows+=t.extract()
            else:
                for line in page.get_text("text").splitlines():
                    if line.strip(): rows.append([no,line.strip()])
    finally: pdf.close()
    pd.DataFrame(rows or [["No extractable data"]]).to_excel(out,index=False,header=False); return out

def docx_pdf(src,out):
    with tempfile.TemporaryDirectory() as td:
        try:
            subprocess.run(["libreoffice","--headless","--convert-to","pdf","--outdir",td,str(src)],check=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE)
            g=Path(td)/(src.stem+".pdf")
            if g.exists(): g.replace(out); return out
        except (FileNotFoundError,subprocess.CalledProcessError): pass
    d=Document(src); c=canvas.Canvas(str(out),pagesize=A4); _,h=A4; y=h-40
    for p in d.paragraphs:
        for chunk in [p.text[i:i+100] for i in range(0,len(p.text),100)]:
            if y<40: c.showPage(); y=h-40
            c.drawString(35,y,chunk); y-=14
    c.save(); return out

def table_pdf(src,out):
    df=pd.read_csv(src) if src.suffix.lower()==".csv" else pd.read_excel(src)
    c=canvas.Canvas(str(out),pagesize=A4); w,h=A4; m=25; y=h-35; cw=(w-2*m)/max(1,len(df.columns))
    def row(vals,bold=False):
        nonlocal y; c.setFont("Helvetica-Bold" if bold else "Helvetica",7)
        for i,v in enumerate(vals): c.drawString(m+i*cw,y,str(v)[:18])
        y-=11
    row(list(df.columns),True)
    for vals in df.astype(str).itertuples(index=False):
        if y<30: c.showPage(); y=h-35; row(list(df.columns),True)
        row(vals)
    c.save(); return out

def pdf_images(src,outdir):
    pdf=fitz.open(src); paths=[]
    try:
        for i,p in enumerate(pdf,1):
            pix=p.get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False); path=outdir/f"page_{i}.jpg"; pix.save(path); paths.append(path)
    finally: pdf.close()
    zpath=outdir/"pdf_images.zip"
    with zipfile.ZipFile(zpath,"w",zipfile.ZIP_DEFLATED) as z:
        for p in paths: z.write(p,p.name)
    return zpath
